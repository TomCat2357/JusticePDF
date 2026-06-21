"""Generate JusticePDFの使い方.docx from scratch.

Run after ``dev/build_manual_screenshots.py`` so screenshots exist
under ``dev/manual_assets/``. Overwrites the existing docx.

Run:
    uv run python dev\\build_manual_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(__file__).resolve().parent / "manual_assets"
OUT_DOCX = ROOT / "JusticePDFの使い方.docx"


# ---------- Style helpers ----------

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_default_font(doc: Document, jp: str = "Yu Gothic UI", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic UI"
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), jp)
    rFonts.set(qn("w:ascii"), jp)
    rFonts.set(qn("w:hAnsi"), jp)
    # Heading styles inherit Normal but force eastAsia explicitly so
    # Word doesn't fall back to a no-CJK font for headings.
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        if h in doc.styles:
            hs = doc.styles[h].element.get_or_add_rPr()
            hr = hs.find(qn("w:rFonts"))
            if hr is None:
                hr = OxmlElement("w:rFonts")
                hs.append(hr)
            hr.set(qn("w:eastAsia"), jp)
            hr.set(qn("w:ascii"), jp)
            hr.set(qn("w:hAnsi"), jp)


def _h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    return p


def _h2(doc, text: str):
    return doc.add_heading(text, level=2)


def _h3(doc, text: str):
    return doc.add_heading(text, level=3)


def _para(doc, text: str = ""):
    return doc.add_paragraph(text)


def _img(doc, name: str, width_cm: float = 13.5):
    path = ASSETS / name
    if not path.exists():
        p = doc.add_paragraph(f"[画像が見つかりません: {name}]")
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def _bullet(doc, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


# 丸数字（①〜⑳）。初学者向けに手順を「①②③…」で示すために使う。
_CIRCLED = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"


def _step(doc, n: int, text: str):
    """番号付きの手順を 1 ステップ追加する（①②③…）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    mark = _CIRCLED[n - 1] if 1 <= n <= len(_CIRCLED) else f"{n}."
    r = p.add_run(mark + " ")
    r.bold = True
    p.add_run(text)
    return p


def _notebox(doc, label: str, text: str, fill_hex: str = "FFF6D6") -> None:
    """色付きの囲み（ヒント・注意など）を 1 つ追加する。"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill_hex)
    p = cell.paragraphs[0]
    r = p.add_run(label + "　")
    r.bold = True
    p.add_run(text)
    _para(doc)  # 後ろに余白


def _tip(doc, text: str) -> None:
    _notebox(doc, "【ヒント】", text, "EAF4FF")


def _note(doc, text: str) -> None:
    _notebox(doc, "【注意】", text, "FFF1F1")


# ---------- Content ----------

# ボタンは実際のツールバー（左→右）の並び順に対応。
# 「▼」付きはクリックでドロップダウンメニューが開くボタン。
BUTTON_TABLE = [
    ("元に戻す", "Ctrl+Z", "直前の操作を取り消します。"),
    ("やり直し", "Ctrl+Y", "取り消した操作をやり直します。"),
    ("削除", "Delete", "選択中の PDF・フォルダをゴミ箱へ移動します。"),
    ("名前変更 ▼", "F2 / Shift+F2",
     "メニューから「ファイル名」（F2）「PDF名」（Shift+F2）「フォルダ名」を選んで変更します。"),
    ("結合", "—",
     "選択した複数のファイル・フォルダを 1 つの PDF に結合します"
     "（フォルダ構成はしおりの階層として再現）。"),
    ("インポート ▼", "—",
     "メニューから「ファイルをインポート」「フォルダをインポート」を選べます"
     "（Office 文書・画像は自動で PDF へ変換）。"),
    ("新規作成 ▼", "—",
     "メニューから「ファイル」（空の 1 ページ PDF）または「フォルダ」（新規サブフォルダ）を作成します。"),
    ("エクスポート", "Ctrl+E",
     "選択した PDF・フォルダを別の場所にコピーで書き出します。形式や圧縮の設定は専用ダイアログから。"),
    ("印刷", "Ctrl+P", "選択した PDF を印刷します。"),
    ("回転", "—", "選択した PDF・ページを時計回りに 90° 回転します。"),
    ("すべて選択", "Ctrl+A", "メイン画面のカード／フォルダを全部選びます。"),
    ("並び替え ▼", "—",
     "メニューから名前順／日付順（各 昇順・降順）に並べ替えます。フォルダカードも対象です。"),
]


def build_doc() -> Document:
    doc = Document()
    _set_default_font(doc)

    # 表紙的な扱いの Heading 1
    _h1(doc, "JusticePDFについて")
    _para(
        doc,
        "JusticePDF（ジャスティス ピーディーエフ）は、PDF ファイルを「カード」のように画面に並べて、"
        "マウスでつかんで動かす（ドラッグ＆ドロップ）だけで、結合・並べ替え・ページの抜き出し・回転などを"
        "直感的に行えるデスクトップアプリです。Word・Excel・PowerPoint・画像ファイルを画面に放り込むと、"
        "自動で PDF に変換して取り込みます。難しい設定は不要で、マウス操作が中心です。",
    )
    _para(doc, "このアプリでできる主なことは次のとおりです。")
    _bullet(doc, "複数の PDF を 1 つにまとめる（結合）")
    _bullet(doc, "ページの順番を入れ替える・不要なページを消す")
    _bullet(doc, "Word・Excel・画像などを PDF に変換して取り込む")
    _bullet(doc, "PDF に付箋・マーカー・図形などの書き込み（注釈）を付ける")
    _bullet(doc, "PDF・画像として保存（エクスポート）したり、印刷したりする")

    _h2(doc, "このマニュアルの読み方")
    _para(
        doc,
        "「２　インストール方法」で導入のしかたを説明し、「３　使い方」で実際の操作を画面の図とともに"
        "順番に解説します。はじめて使う方は、最初から順に読み進めてください。"
        "文中の「クリック」はマウスの左ボタンを 1 回押すこと、「ダブルクリック」は左ボタンを"
        "すばやく 2 回押すこと、「ドラッグ＆ドロップ」は左ボタンを押したままマウスを動かし、"
        "目的の場所で指を離すことを指します。",
    )

    _h2(doc, "用語の説明（はじめての方へ）")
    _para(doc, "本マニュアルでよく出てくる言葉を、先にまとめておきます。")
    glossary = [
        ("PDF（ピーディーエフ）",
         "どのパソコンでも見た目が崩れずに表示できる、文書の保存形式です。"),
        ("カード",
         "メイン画面で 1 つの PDF を表す四角い絵（サムネイル）のことです。"
         "トランプのカードのように並べて扱えます。"),
        ("フォルダカード",
         "メイン画面で、PDF をまとめて入れておく「フォルダ」を表すカードです。"),
        ("ドラッグ＆ドロップ",
         "マウスの左ボタンを押したまま動かして（ドラッグ）、目的の場所で離す（ドロップ）操作です。"),
        ("注釈（ちゅうしゃく）",
         "PDF に後から書き加える、付箋・マーカー・図形・コメントなどの書き込みのことです。"),
        ("しおり",
         "PDF の中の見出し（目次）です。クリックすると、その場所へジャンプできます。"),
    ]
    gtable = doc.add_table(rows=1, cols=2)
    gtable.style = "Light Grid Accent 1"
    ghdr = gtable.rows[0].cells
    ghdr[0].text = "用語"
    ghdr[1].text = "意味"
    for c in ghdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E5FF")
    for term, desc in glossary:
        row = gtable.add_row().cells
        row[0].text = term
        row[1].text = desc
    for row in gtable.rows:
        for c, w in zip(row.cells, (Cm(4.5), Cm(11.5))):
            c.width = w
    _para(doc)

    # 1. 利用規約
    _h1(doc, "１　利用規約")
    _para(doc, "以下の利用規約を復唱し、同意してください。※難しければ心の中で。")
    _para(doc, "規約")

    # 2. インストール方法（現行を流用）
    _h1(doc, "２　インストール方法")
    _para(
        doc,
        "JusticePDF は「ZIP（ジップ）ファイル」という、圧縮された 1 つのまとまりで配られます。"
        "これをパソコンに保存し、中身を取り出す（展開する）と使えるようになります。",
    )

    _h2(doc, "STEP 1　ダウンロード")
    _step(doc, 1, "JusticePDF のフォルダ内にある「JusticePDF.zip」を見つけます。")
    _step(doc, 2, "それを、自分のパソコンのわかりやすい場所（例：ドキュメント）に保存します。")
    _img(doc, "16_install_zip.png")

    _h2(doc, "STEP 2　ZIP ファイルの展開（中身を取り出す）")
    _step(doc, 1, "保存した「JusticePDF.zip」を右クリックします。")
    _step(doc, 2, "出てきたメニューから「すべて展開」を選びます。")
    _img(doc, "17_install_extract_menu.png")

    _step(doc, 3, "確認の画面が出たら、「展開」ボタンをクリックします。")
    _img(doc, "19_install_extract_button.png")
    _note(doc, "展開には数分かかることがあります。終わるまで待ってください。"
          "終わると、中身が入ったフォルダが開きます。")

    _h2(doc, "STEP 3（任意）　Windows 統合設定（setup_all）")
    _tip(
        doc,
        "ここは「やっておくと便利になる」設定です。とりあえず使い始めたい方は、"
        "この STEP 3 を飛ばして「３　使い方」へ進んでも構いません。",
    )
    _para(
        doc,
        "ZIP を展開しただけでも「run_justice_gui.cmd」から起動できますが、"
        "以下の統合設定を行うと Windows への組み込みがより快適になります。"
        "すべてユーザー単位の設定（HKCU）で、管理者権限は不要です。",
    )
    _para(
        doc,
        "展開したフォルダをエクスプローラーで開き、アドレスバーに"
        " PowerShell と入力して Enter を押します（または「ファイル」タブ →「Windows PowerShell を開く」）。"
        "表示された PowerShell ウィンドウで次のコマンドを実行してください：",
    )

    # コードブロック風の段落
    code_para = doc.add_paragraph()
    code_para.style = "No Spacing"
    run = code_para.add_run(
        "powershell -ExecutionPolicy Bypass -File tools\\setup_all.ps1"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    _para(doc)  # spacer

    _para(
        doc,
        "このコマンドは以下の 2 つのステップを順番に実行します：",
    )

    SETUP_STEPS = [
        ("ステップ 1：右クリックメニューの追加",
         "エクスプローラーで PDF・Office 文書・画像ファイル・フォルダを右クリックすると"
         "「JusticePDFで開く」が現れるようになります（Windows 11 では「その他のオプション」内）。"),
        ("ステップ 2：デスクトップショートカットの作成",
         "デスクトップに「JusticePDF」ショートカットを作成します（アイコンを自動生成）。"
         "グローバルホットキー Ctrl+Alt+J が割り当てられ、どこからでもJusticePDFを起動できます。"),
    ]

    for title, desc in SETUP_STEPS:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title + "：")
        r.bold = True
        p.add_run(desc)

    _para(doc)

    _para(doc, "【主なオプション】")

    SETUP_OPTIONS = [
        ("-StartMenu",
         "デスクトップに加え、スタートメニューにもショートカットを作成します。"
         "タスクバーへのピン留めが容易になります。"),
        ("-Background",
         "フォルダ内の何もない場所を右クリックしたときにも「JusticePDFで開く」を表示します。"),
        ("-ClassicMenu",
         "Windows 11 のクラシックコンテキストメニューを有効化し、"
         "「その他のオプション」を開かなくてもトップレベルに「JusticePDFで開く」を表示します。"
         "適用には Explorer の再起動が必要です。"),
        ("-SkipShortcut",
         "ステップ 2（デスクトップショートカット）をスキップします。"),
        ("-Hotkey \"Ctrl+Alt+J\"",
         "ショートカットのホットキーを変更します（デフォルトは Ctrl+Alt+J）。"),
        ("-ContinueOnError",
         "いずれかのステップが失敗しても残りのステップを続行します。"),
    ]

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "オプション"
    hdr2[1].text = "説明"
    for c in hdr2:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E5FF")
    for opt, desc in SETUP_OPTIONS:
        row = table2.add_row().cells
        row[0].text = opt
        row[1].text = desc
    widths2 = [Cm(4.5), Cm(11.5)]
    for row in table2.rows:
        for c, w in zip(row.cells, widths2):
            c.width = w

    _para(doc)

    _para(
        doc,
        "例：右クリックにクラシックメニューでトップ表示＋スタートメニューにも登録する場合：",
    )
    code_para2 = doc.add_paragraph()
    code_para2.style = "No Spacing"
    run2 = code_para2.add_run(
        "powershell -ExecutionPolicy Bypass -File tools\\setup_all.ps1 "
        "-Background -ClassicMenu -StartMenu"
    )
    run2.font.name = "Courier New"
    run2.font.size = Pt(10)
    _para(doc)

    _para(
        doc,
        "※ タスクバーへのピン留めは Windows の制限により自動化できません。"
        "デスクトップの「JusticePDF」ショートカットを右クリック →「タスクバーにピン留めする」"
        "（Windows 11 では「その他のオプションを表示」→「タスクバーにピン留めする」）で手動で行ってください。",
    )

    # 3. 使い方
    _h1(doc, "３　使い方")

    _h2(doc, "3.1　起動と作業フォルダ")
    _para(doc, "まずはアプリを起動してみましょう。")
    _step(doc, 1, "展開したフォルダの中にある「run_justice_gui.cmd」をダブルクリックします。")
    _step(doc, 2, "「実行しますか？」のような確認画面が出たら、「実行」をクリックします。")
    _step(doc, 3, "しばらくすると JusticePDF の画面が開きます。")
    _para(
        doc,
        "起動すると、最初に「ドキュメント／PDFs」というフォルダが"
        "「作業フォルダ」（＝今あつかっているフォルダ）として開きます。"
        "起動直後は、他のウィンドウより手前に表示されます。",
    )
    _tip(
        doc,
        "STEP 3 でデスクトップショートカットを作った場合は、"
        "デスクトップの「JusticePDF」アイコンをダブルクリックしても起動できます。",
    )

    _h2(doc, "3.2　ファイルの取り込み")
    _para(
        doc,
        "JusticePDF であつかいたいファイルを、画面に取り込みます。"
        "もっとも簡単なのは、ファイルをマウスでつかんで画面に放り込む方法です。",
    )
    _para(doc, "【方法 A：ドラッグ＆ドロップで放り込む】")
    _step(doc, 1, "エクスプローラーなどで、取り込みたいファイルを選びます。")
    _step(doc, 2, "そのファイルをマウスでつかんだまま、JusticePDF の画面の上まで動かし、指を離します。")
    _para(
        doc,
        "PDF はもちろん、Word（.doc/.docx）・Excel（.xls/.xlsx）・PowerPoint（.ppt/.pptx）・"
        "画像（.png/.jpg/.bmp/.tif/.gif）も取り込めます。Office 文書や画像は、取り込むときに"
        "自動で PDF に変換されます。",
    )
    _para(doc, "【方法 B：「インポート」ボタンから選ぶ】")
    _para(
        doc,
        "画面上部の「インポート」ボタンをクリックすると、下の図のようなメニューが出ます。",
    )
    _img(doc, "20_import_menu.png")
    _bullet(doc, "「ファイルをインポート」：ファイルを 1 つずつ選んで取り込みます。")
    _bullet(doc, "「フォルダをインポート」：フォルダごと、中身をまとめて取り込みます。")

    _h2(doc, "3.3　メイン画面の見方")
    _para(
        doc,
        "アプリのメインとなる画面です。ここに、取り込んだ PDF が「カード」として、"
        "フォルダが「フォルダカード」として並びます。",
    )
    _bullet(doc, "PDF のカード：右上の数字（例：4p）は、その PDF のページ数です。")
    _bullet(doc, "フォルダカード：右上の数字は、その中に入っている項目の数です。")
    _bullet(doc, "画面のいちばん上にある横長の帯が「ツールバー」で、操作ボタンが並んでいます。")
    _img(doc, "01_main_idle.png")

    _h2(doc, "3.4　ボタン機能一覧")
    _para(
        doc,
        "メイン画面の上部にあるツールバーには、次のボタンが左から順に並んでいます。"
        "各ボタンの役割は下の表のとおりです。ボタン名の横に「▼」が付いているものは、"
        "クリックすると下にメニューが開いて、さらに細かく選べます。",
    )
    _img(doc, "04_toolbar.png")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "ボタン"
    hdr[1].text = "ショートカット"
    hdr[2].text = "機能"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E5FF")
    for name, shortcut, desc in BUTTON_TABLE:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = shortcut
        row[2].text = desc
    # Column widths
    widths = [Cm(3.0), Cm(2.6), Cm(10.5)]
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = w

    _para(doc)  # spacer

    _para(doc, "「▼」付きボタンの例として、「名前変更」と「並び替え」のメニューを示します。")
    _img(doc, "23_rename_menu.png")
    _img(doc, "22_sort_menu.png")
    _tip(
        doc,
        "「ショートカット」の欄は、キーボードで同じ操作をするための「近道キー」です。"
        "たとえば「Ctrl+Z」は、Ctrl キーを押しながら Z キーを押すという意味です。"
        "覚えなくても、すべてボタンから操作できます。",
    )

    _h2(doc, "3.5　選択操作（カードの選び方）")
    _para(
        doc,
        "ファイルを操作する前に、まず対象のカードを「選択」します。"
        "選択されたカードは、青いわくで囲まれて目立つようになります。",
    )
    _bullet(doc, "クリック：1 枚だけ選ぶ（単一選択）")
    _bullet(doc, "Ctrl を押しながらクリック：選択を 1 枚ずつ追加・取り消し（複数選択）")
    _bullet(doc, "Shift を押しながらクリック：最初に選んだカードから今のカードまでをまとめて選ぶ（範囲選択）")
    _bullet(doc, "何もない場所からドラッグ：四角い枠で囲んでまとめて選ぶ（フォルダカードも一緒に選べます）")
    _img(doc, "02_main_selection.png")

    _h2(doc, "3.6　並べ替え・結合")
    _para(
        doc,
        "カードをマウスでつかんで動かす（ドラッグ）と、置く場所によって 2 つの結果になります。"
        "カードとカードの「あいだ」に置くと並び順が変わり、カードの「真ん中」に重ねるように置くと"
        "2 つの PDF が 1 つに合体します（結合）。結合できるときは、相手のカードが緑色に光って教えてくれます。",
    )
    _bullet(doc, "カード間にドロップ：並べ替え")
    _bullet(doc, "Ctrl を押しながらカード間：複製して挿入（_copy_N.pdf を作成）")
    _bullet(doc, "カード中央へドロップ：結合（元のカードに合体、移動元はゴミ箱へ）")
    _bullet(doc, "Ctrl を押しながら中央へドロップ：複製して結合（移動元は残ります）")
    _img(doc, "03_main_drop_merge.png")
    _para(
        doc,
        "ドラッグ＆ドロップでカードを重ねて結合した場合、重ねた各 PDF が元々持っている"
        "しおりは、章・節などの親子関係（階層）をそのまま保って引き継がれます。"
        "ファイル名のしおりは付かないため、1 つずつ繰り返し重ねてもしおりが"
        "余計に深くなることはありません。",
    )
    _para(
        doc,
        "複数のファイル・フォルダを選んでツールバーの「結合」ボタンを押すと、"
        "選択した項目をまとめて 1 つの PDF に結合できます。フォルダを含めて結合すると、"
        "フォルダ構成がしおり（目次）の階層として再現されます。"
        "結合後、元のファイル・フォルダはゴミ箱へ移動します（「元に戻す」で復元できます）。",
    )

    _h2(doc, "3.7　新規作成（空の PDF・フォルダ）")
    _para(
        doc,
        "ツールバーの「新規作成」ボタンをクリックすると、下の図のメニューが開きます。"
        "空のファイルや、整理用のフォルダをその場で作れます。",
    )
    _img(doc, "21_new_menu.png")
    _bullet(doc, "「ファイル」：空っぽの 1 ページの PDF（新規ファイル.pdf）を作ります。")
    _bullet(doc, "「フォルダ」：新しいサブフォルダ（新規フォルダ）を作ります。")
    _para(
        doc,
        "作った後の名前は「名前変更」ボタンで変えられます。"
        "うっかり作ってしまっても、「元に戻す」ボタン（または Ctrl+Z）で取り消せます。",
    )

    _h2(doc, "3.8　フォルダの使い方")
    _para(
        doc,
        "フォルダカードをダブルクリックすると、そのフォルダを別ウィンドウとして開けます。"
        "フォルダ同士の並べ替え・移動、複数選択してまとめて削除も可能です。"
        "「並び替え」ボタンはフォルダカードにも適用されます。"
        "フォルダの中身（PDF の追加・削除）が外で変わると、自動で更新されます。",
    )
    _img(doc, "11_folder_selection.png")

    _h2(doc, "3.9　個別画面（ページ編集）")
    _para(
        doc,
        "1 つの PDF の「中身（ページ）」を編集したいときは、メイン画面でそのカードを"
        "ダブルクリックします。すると、その PDF のページが 1 枚ずつ並んだ"
        "「個別画面」が開きます。ここでは、ページの並べ替え・削除・回転や、"
        "別の PDF からのページの差し込みができます。"
        "ページをメイン画面の方へドラッグすれば、そのページだけを取り出して新しい PDF にできます。",
    )
    _img(doc, "05_page_edit_grid.png")
    _para(doc, "ページに対してできる主なドラッグ操作は次のとおりです。")
    _bullet(doc, "ページ間にドロップ：同一 PDF 内で並べ替え")
    _bullet(doc, "他 PDF のページをドロップ：その位置にページ挿入")
    _bullet(doc, "ページをメイン画面のカードへドロップ：そのカードの先頭にページ挿入")
    _bullet(doc, "ページをメイン画面の空き場所へドロップ：抽出ページで新しい PDF を作成")
    _bullet(doc, "Ctrl 押下時はコピー、無しはムーブ（移動元から削除）")
    _bullet(doc, "メイン画面のカードを個別画面にドロップ：その PDF を丸ごとページとして挿入")

    _h2(doc, "3.10　拡大ビュー（大きく表示して読む）")
    _para(
        doc,
        "ページを大きく表示してじっくり読みたいときは、個別画面でページの絵（サムネイル）を"
        "ダブルクリックします。すると「拡大ビュー」に切り替わり、ページが大きく表示されます。"
        "マウスのホイールを回すと拡大・縮小でき、文字をなぞって選択しコピーしたり、"
        "ページ内のリンクをクリックしたりできます。文字は 1 文字単位で選べ、"
        "行をまたいでドラッグしても自然に範囲を選択できます。",
    )
    _bullet(doc, "PageUp / PageDown：前後のページへ")
    _bullet(doc, "Home / End：先頭・末尾のページへ")
    _bullet(doc, "「100%」ボタン：クリックすると 25%〜400% の倍率プリセットをドロップダウンで選択")
    _bullet(doc, "矢印キー：表示位置をスクロール。Ctrl + 矢印で大きく移動")
    _bullet(doc, "マウス中ボタンでドラッグ：紙をつかむように表示位置を移動（パン）")
    _bullet(doc, "マウス右ボタンでドラッグ：囲んだ範囲をビュー全体に拡大表示")
    _bullet(doc, "Back ボタン：拡大ビューを抜けて元の一覧に戻る")
    _img(doc, "06_zoom_view.png")

    _h3(doc, "見開き表示")
    _para(
        doc,
        "本のように 2 ページを左右に並べて読みたいときは、「見開き表示」ボタンを押します。",
    )
    _img(doc, "25_spread_view.png")
    _para(
        doc,
        "見開き表示は「閲覧専用（読むだけ）」モードです。そのため、文字選択・注釈編集・"
        "リンク操作はできません。ホイールでのズーム、マウス中ボタンでのパン（移動）、"
        "右ドラッグでの範囲拡大、矢印キーでのスクロール（Ctrl で大きく移動）はそのまま使えます。"
        "前後への移動は 2 ページずつ進みます。もう一度ボタンを押すと、元の 1 ページ表示に戻ります"
        "（拡大ビューを開き直したときも、最初は 1 ページ表示で始まります）。",
    )

    _h2(doc, "3.11　注釈（書き込み）の追加")
    _para(
        doc,
        "「注釈（ちゅうしゃく）」とは、PDF の上に後から書き加える、付箋・マーカー・図形・"
        "コメントなどの書き込みのことです。元の文章はそのままに、目印やメモを重ねられます。",
    )
    _para(
        doc,
        "注釈を付けるには、まず拡大ビューにします（3.10 参照）。次に画面の右はしにある"
        "「アノテーション」ボタン、または右端の「▶」ボタンを押すと、下の図のような"
        "「注釈パネル」が右側に開きます。",
    )
    _img(doc, "07_annotation_panel.png")
    _para(
        doc,
        "パネルの「新規」ボタンで付箋（文字を書き込める箱）を追加でき、"
        "図形ボタン（―、△、○、□、[ ]）で線・三角・楕円・四角・かっこを描けます。"
        "幅・高さ・線の太さ・透明度・文字サイズ・文字色・背景色・線の色を、細かく指定できます。",
    )
    _bullet(doc, "注釈を選んで矢印キー：移動（通常10pt、Alt/Shiftで微調整1pt、Ctrlで粗く50pt）")
    _bullet(doc, "注釈を選んで Delete：削除")
    _bullet(doc, "Ctrl + ドラッグ：注釈を複製")
    _bullet(doc, "Ctrl + C / Ctrl + V：別ページへコピー＆ペースト")
    _bullet(doc, "「最背面」「背面へ」「前面へ」「最前面」：重なり順を変更")
    _bullet(doc, "線注釈は端点ハンドルでドラッグ可。Shift 押下で水平・垂直・45° にスナップ")

    _h3(doc, "テキストマーカー（ハイライト・下線・取り消し線）")
    _para(
        doc,
        "重要な文に色を付けたり、線を引いたりできます。",
    )
    _step(doc, 1, "ページ上の文字を、マウスでなぞって選択します。")
    _step(doc, 2, "注釈パネルの「マーカー」「U」「S」ボタンのいずれかを押します。")
    _para(
        doc,
        "「マーカー」はハイライト（蛍光ペン）、「U」は下線（Underline）、"
        "「S」は取り消し線（Strikeout）です。下の図は、3 種類を付けた例です。",
    )
    _img(doc, "27_markup.png")
    _para(
        doc,
        "作ったマーカーはクリックで選べ、「マークアップの色」ボタンで色を変えたり、"
        "別の種類に切り替えたり、Delete キーで消したりできます"
        "（すべて「元に戻す」でやり直せます）。",
    )

    _h3(doc, "付箋ノート（コメント）")
    _para(
        doc,
        "ページの好きな位置に、コメントを書き込める小さな付箋アイコンを置けます。",
    )
    _step(doc, 1, "注釈パネルの「ノート」ボタンを押します。")
    _step(doc, 2, "ページ上の、付箋を置きたい場所をクリックします。")
    _para(
        doc,
        "置いた付箋アイコンにマウスを乗せると、下の図のように中身がふきだしで表示されます。"
        "クリックして選ぶと、本文の編集や「付箋の色」の変更ができます。"
        "注釈パネルの下のほうには、そのページの付箋の一覧も表示されます。",
    )
    _img(doc, "28_note.png")

    _h3(doc, "校正コールアウト（修正の指示）")
    _para(
        doc,
        "「ここをこう直して」という校正の指示を、矢印付きで書き込めます。",
    )
    _step(doc, 1, "注釈パネルの「校正」ボタンを押します。")
    _step(doc, 2, "直したい位置をクリックします。")
    _para(
        doc,
        "すると、その位置を指す矢印（引き出し線）と、修正内容を書き込む文字ボックスが"
        "ひとまとまりで置かれ、すぐに文字を入力できます。"
        "矢印の先（さきっぽ）は、後からドラッグして指す位置を微調整できます。"
        "まとまりのどれかを消すと校正コールアウト全体が消え、「元に戻す」で戻せます。",
    )
    _img(doc, "29_callout.png")

    _h3(doc, "しおり（目次）パネル")
    _para(
        doc,
        "拡大ビューで「しおり」ボタンを押すと、右側に「しおり（＝PDF の目次）」の"
        "一覧が表示されます。見出しをクリックすると、その場所へすぐに移動できます。"
        "付けた付箋は、直前のしおりごとに「付箋(N)」としてまとめられ、"
        "クリックするとその付箋のあるページへジャンプします。",
    )
    _img(doc, "26_bookmarks_panel.png")

    _h2(doc, "3.12　ページ内検索（Ctrl + F）")
    _para(
        doc,
        "PDF の中から目的の言葉をさがせます。",
    )
    _step(doc, 1, "個別画面（または拡大ビュー）で、キーボードの Ctrl キーを押しながら F キーを押します。")
    _step(doc, 2, "出てきた小さな入力欄に、さがしたい言葉を入力して Enter キーを押します。")
    _step(doc, 3, "「次へ」「前へ」のボタンで、見つかった場所を順に移動します。")
    _para(doc, "見つかった文字は、ページ上で色が付いて目立つように表示されます。")
    _img(doc, "09_search_dialog.png", width_cm=10)

    _h2(doc, "3.13　印刷")
    _para(
        doc,
        "印刷したい PDF のカードを選んでから、「印刷」ボタン（または Ctrl+P）を押すと、"
        "下の図の印刷画面が開きます。左側で設定をすると、右側の見本（プレビュー）に"
        "仕上がりがその場で表示されます。",
    )
    _img(doc, "24_print_dialog.png", width_cm=14)
    _bullet(doc, "プリンタ：印刷に使う機械を選びます。")
    _bullet(doc, "部数：何部印刷するかを指定します。")
    _bullet(doc, "印刷範囲：すべて／現在のページ／ページ指定（例：1-5, 8）から選べます。")
    _bullet(doc, "用紙サイズ・向き・両面・カラー／モノクロ・1 枚あたりのページ数なども指定できます。")
    _step(doc, 1, "設定を確認します（右の見本で仕上がりをチェックできます）。")
    _step(doc, 2, "問題なければ、右下の「印刷」ボタンを押します。")

    _h2(doc, "3.14　エクスポート（別の場所へ保存）")
    _para(
        doc,
        "「エクスポート」とは、編集した PDF を別のファイルとして外に書き出す（保存する）ことです。"
        "「エクスポート」ボタン（または Ctrl+E）を押すと、下の図の「エクスポート設定」画面が開きます。",
    )
    _para(
        doc,
        "PDF・PNG・JPEG の 3 つの形式で保存できます（PNG・JPEG は画像形式です）。"
        "PDF で保存するときは、ファイルを軽くするための圧縮レベル"
        "（なし／軽量／高圧縮／強圧縮／中強圧縮／より強い圧縮／かなり強い圧縮／最大圧縮／カスタム）と、"
        "「テキストデータを削除（画像のみ）」を選べます。"
        "PNG・JPEG では、画像のきめ細かさ（DPI、72〜600）と JPEG の画質を指定できます。",
    )
    _para(
        doc,
        "フォルダカードを選んでエクスポートすると、サブフォルダの構造を保ったまま"
        "中の PDF をまとめて書き出します（PDF 以外のファイルは対象外）。"
        "ファイルとフォルダを同時に選んでまとめてエクスポートすることもできます。",
    )
    _img(doc, "10_export_dialog.png", width_cm=10)

    _h2(doc, "3.15　ファイルを外に出す")
    _para(
        doc,
        "メイン画面のカードを、JusticePDF の外（デスクトップ・エクスプローラー・他のアプリなど）へ"
        "ドラッグ＆ドロップすると、ファイルとしてそのまま取り出せます。"
        "また、カードを右クリックすると出るメニューから、「印刷」や「他のアプリで開く」も選べます。",
    )

    _h2(doc, "3.16　マルチウィンドウ（複数の窓を並べる）")
    _para(
        doc,
        "フォルダカードをダブルクリックすると、そのフォルダを別の窓（ウィンドウ）として開けます。"
        "下の図のように 2 つの作業フォルダを左右に並べ、ドラッグ＆ドロップで"
        "ファイルやフォルダをやり取りできます。整理に便利です。",
    )
    _img(doc, "12_multi_window.png")

    _h2(doc, "3.17　困ったときは（操作をなかったことにする）")
    _para(
        doc,
        "操作をまちがえても、あわてる必要はありません。",
    )
    _bullet(doc, "「元に戻す」ボタン（Ctrl+Z）：直前の操作を取り消します。何回でもさかのぼれます。")
    _bullet(doc, "「やり直し」ボタン（Ctrl+Y）：取り消した操作を、もう一度やり直します。")
    _para(
        doc,
        "削除したファイルやフォルダは、すぐに消えるのではなく「ゴミ箱」へ移動します。"
        "「元に戻す」で復元できますので、安心して操作してください。",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
