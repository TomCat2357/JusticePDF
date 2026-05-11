"""Generate JusticePDFの使い方.docx from scratch.

Run after ``tools/build_manual_screenshots.py`` so screenshots exist
under ``tools/manual_assets/``. Overwrites the existing docx.

Run:
    uv run python tools\\build_manual_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "tools" / "manual_assets"
OUT_DOCX = ROOT / "JusticePDFの使い方.docx"


# ---------- Style helpers ----------

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_default_font(doc: Document, jp: str = "Yu Gothic UI", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic UI"
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), jp)
    rFonts.set(qn("w:ascii"), jp)
    rFonts.set(qn("w:hAnsi"), jp)
    # Heading styles inherit Normal but force eastAsia explicitly so
    # Word doesn't fall back to a no-CJK font for headings.
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        if h in doc.styles:
            hs = doc.styles[h].element.get_or_add_rPr()
            hr = hs.find(qn("w:rFonts"))
            if hr is None:
                hr = OxmlElement("w:rFonts")
                hs.append(hr)
            hr.set(qn("w:eastAsia"), jp)
            hr.set(qn("w:ascii"), jp)
            hr.set(qn("w:hAnsi"), jp)


def _h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    return p


def _h2(doc, text: str):
    return doc.add_heading(text, level=2)


def _h3(doc, text: str):
    return doc.add_heading(text, level=3)


def _para(doc, text: str = ""):
    return doc.add_paragraph(text)


def _img(doc, name: str, width_cm: float = 13.5):
    path = ASSETS / name
    if not path.exists():
        p = doc.add_paragraph(f"[画像が見つかりません: {name}]")
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def _bullet(doc, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


# ---------- Content ----------

BUTTON_TABLE = [
    ("Undo", "Ctrl+Z", "直前の操作を取り消します。"),
    ("Redo", "Ctrl+Y", "取り消した操作をやり直します。"),
    ("Delete", "Delete", "選択中の PDF・フォルダをゴミ箱へ移動します。"),
    ("Rename", "F2", "選択した PDF のファイル名を変更します（1 件選択時のみ）。"),
    ("Title", "Shift+F2", "PDF のメタデータタイトルを確認・変更します。"),
    ("Import", "—", "PDF・Word・Excel・PowerPoint・画像をメイン画面へ取り込みます（自動で PDF へ変換）。"),
    ("Import Folder", "—", "サブフォルダごと一括で取り込みます。"),
    ("New Folder", "—", "作業フォルダ内に新しいサブフォルダを作成します。"),
    ("Export", "Ctrl+E", "選択した PDF（または全部）を別の場所にコピーで書き出します。形式や圧縮の設定は専用ダイアログから。"),
    ("Print", "Ctrl+P", "選択した PDF を印刷します。"),
    ("Refresh", "—", "ファイルの実体と画面の表示を同期します（外部で書き換えた時に）。"),
    ("Rotate", "—", "選択した PDF・ページを時計回りに 90° 回転します。"),
    ("Select All", "Ctrl+A", "メイン画面のカード／フォルダを全部選びます。"),
    ("Sort Name", "—", "名前順に並べ替えます（もう一度押すと逆順）。"),
    ("Sort Date", "—", "更新日時順に並べ替えます（もう一度押すと逆順）。"),
]


def build_doc() -> Document:
    doc = Document()
    _set_default_font(doc)

    # 表紙的な扱いの Heading 1
    _h1(doc, "JusticePDFについて")
    _para(
        doc,
        "JusticePDF は、PDF をカードのように並べて、ドラッグ＆ドロップで結合・並べ替え・"
        "ページの抜き出し・回転などを直感的にできるデスクトップアプリです。"
        "Word・Excel・PowerPoint・画像ファイルを放り込むと、自動で PDF に変換して取り込みます。",
    )

    # 1. 利用規約
    _h1(doc, "１　利用規約")
    _para(doc, "以下の利用規約を復唱し、同意してください。※難しければ心の中で。")
    _para(doc, "規約")

    # 2. インストール方法（現行を流用）
    _h1(doc, "２　インストール方法")
    _h2(doc, "ダウンロード")
    _para(doc, "JusticePDF のフォルダ内にある JusticePDF.zip を、ダウンロードします。")
    _img(doc, "16_install_zip.png")

    _h2(doc, "ZIP ファイル展開")
    _para(doc, "（この前に Document など好きな場所に zip ファイルを移動してもいいです）")
    _para(doc, "右クリック → 「すべて展開」 を選びます。")
    _img(doc, "17_install_extract_menu.png")

    _para(doc, "「展開」ボタンをクリックします。")
    _img(doc, "19_install_extract_button.png")
    _para(doc, "※展開には数分かかります。")

    # 3. 使い方
    _h1(doc, "３　使い方")

    _h2(doc, "3.1　起動と作業フォルダ")
    _para(
        doc,
        "展開されたフォルダの中にある「run_justice_gui.cmd」をダブルクリックすると、"
        "JusticePDF が立ち上がります。「実行」を確認するダイアログが出たらクリックしてください。",
    )
    _para(
        doc,
        "起動すると「ドキュメント／PDFs」フォルダが作業フォルダとして開きます。"
        "起動直後は他のウィンドウより前面に表示されます。",
    )

    _h2(doc, "3.2　ファイルの取り込み")
    _para(
        doc,
        "JusticePDF のウィンドウへ、PDF・Word（.doc/.docx）・Excel（.xls/.xlsx）・"
        "PowerPoint（.ppt/.pptx）・画像（.png/.jpg/.bmp/.tif/.gif）を"
        "ドラッグ＆ドロップで放り込むだけで取り込めます。Office 文書や画像は自動で PDF に変換します。",
    )
    _para(doc, "Import ボタンからファイルを選んで取り込むこともできます。")

    _h2(doc, "3.3　メイン画面の見方")
    _para(
        doc,
        "メイン画面は、PDF を「カード」、サブフォルダを「フォルダカード」として並べて表示します。"
        "カードの右上はページ数、フォルダカードの右上はその中の項目数です。",
    )
    _img(doc, "01_main_idle.png")

    _h2(doc, "3.4　ボタン機能一覧")
    _img(doc, "04_toolbar.png")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "ボタン"
    hdr[1].text = "ショートカット"
    hdr[2].text = "機能"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E5FF")
    for name, shortcut, desc in BUTTON_TABLE:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = shortcut
        row[2].text = desc
    # Column widths
    widths = [Cm(3.0), Cm(2.6), Cm(10.5)]
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = w

    _para(doc)  # spacer

    _h2(doc, "3.5　選択操作")
    _bullet(doc, "クリック：単一選択")
    _bullet(doc, "Ctrl + クリック：選択を追加・解除（複数選択）")
    _bullet(doc, "Shift + クリック：範囲選択")
    _bullet(doc, "何もない場所をドラッグ：矩形（ラバーバンド）選択。フォルダカードもまとめて選べます")
    _img(doc, "02_main_selection.png")

    _h2(doc, "3.6　並べ替え・結合")
    _para(
        doc,
        "カードをカードの間にドロップすると順番が変わります。カードの中央にドロップすると "
        "「ファイルの結合」になり、対象のカードが緑色にハイライトされます。",
    )
    _bullet(doc, "カード間にドロップ：並べ替え")
    _bullet(doc, "Ctrl を押しながらカード間：複製して挿入（_copy_N.pdf を作成）")
    _bullet(doc, "カード中央へドロップ：結合（元のカードに合体、移動元はゴミ箱へ）")
    _bullet(doc, "Ctrl を押しながら中央へドロップ：複製して結合（移動元は残ります）")
    _img(doc, "03_main_drop_merge.png")

    _h2(doc, "3.7　フォルダの使い方")
    _para(
        doc,
        "フォルダカードをダブルクリックすると、そのフォルダを別ウィンドウとして開けます。"
        "フォルダ同士の並べ替え・移動、複数選択してまとめて削除も可能です。"
        "フォルダの中身（PDF の追加・削除）が外で変わると、自動で更新されます。",
    )
    _img(doc, "11_folder_selection.png")

    _h2(doc, "3.8　個別画面（ページ編集）")
    _para(
        doc,
        "メイン画面のカードをダブルクリックすると、その PDF のページを並べた個別画面が開きます。"
        "ページの並べ替え、削除、回転、別 PDF からの差し込みができます。"
        "ページをメイン画面側へドラッグすれば、そのページだけを切り出して新しい PDF にできます。",
    )
    _img(doc, "05_page_edit_grid.png")
    _bullet(doc, "ページ間にドロップ：同一 PDF 内で並べ替え")
    _bullet(doc, "他 PDF のページをドロップ：その位置にページ挿入")
    _bullet(doc, "ページをメイン画面のカードへドロップ：そのカードの先頭にページ挿入")
    _bullet(doc, "ページをメイン画面の空き場所へドロップ：抽出ページで新しい PDF を作成")
    _bullet(doc, "Ctrl 押下時はコピー、無しはムーブ（移動元から削除）")
    _bullet(doc, "メイン画面のカードを個別画面にドロップ：その PDF を丸ごとページとして挿入")

    _h2(doc, "3.9　拡大ビュー")
    _para(
        doc,
        "個別画面でページのサムネイルをダブルクリックすると、拡大ビューに切り替わります。"
        "マウスホイール（または Ctrl + ホイール）でズーム、テキストを直接選択してコピー、"
        "リンクのクリックもできます。",
    )
    _bullet(doc, "PageUp / PageDown：前後のページへ")
    _bullet(doc, "Home / End：先頭・末尾のページへ")
    _bullet(doc, "Back ボタン：拡大ビューを抜けて元の一覧に戻る")
    _img(doc, "06_zoom_view.png")

    _h2(doc, "3.10　注釈（付箋・図形）の追加")
    _para(
        doc,
        "拡大ビュー右の「▶」ボタンを押すと、注釈パネルが開きます。"
        "「新規」で付箋（FreeText）を追加し、図形ボタン（―、△、○、□、[ ]）で線・三角・楕円・四角・括弧を描けます。"
        "幅・高さ・線幅・透明度・文字サイズ・文字色・背景色・線色を細かく指定できます。",
    )
    _img(doc, "07_annotation_panel.png")
    _bullet(doc, "注釈を選んで Delete：削除")
    _bullet(doc, "Ctrl + ドラッグ：注釈を複製")
    _bullet(doc, "Ctrl + C / Ctrl + V：別ページへコピー＆ペースト")
    _bullet(doc, "「最背面」「背面へ」「前面へ」「最前面」：重なり順を変更")
    _bullet(doc, "線注釈は端点ハンドルでドラッグ可。Shift 押下で水平・垂直・45° にスナップ")

    _h2(doc, "3.11　ページ内検索（Ctrl + F）")
    _para(
        doc,
        "個別画面で Ctrl + F を押すと検索ダイアログが開きます。"
        "語句を入力して Enter で検索、「次へ」「前へ」でヒット箇所を移動できます。"
        "ヒットした文字はページ上で強調表示されます。",
    )
    _img(doc, "09_search_dialog.png", width_cm=10)

    _h2(doc, "3.12　印刷（Ctrl + P）")
    _para(
        doc,
        "Ctrl + P または Print ボタンで、選択中の PDF を印刷できます。"
        "通常の Windows の印刷ダイアログが開くので、プリンターや範囲を指定して印刷してください。",
    )

    _h2(doc, "3.13　エクスポート（保存）")
    _para(
        doc,
        "Export ボタンまたは Ctrl + E で「エクスポート設定」ダイアログが開きます。"
        "PDF・PNG・JPEG の 3 形式で書き出せます。PDF 出力時は最適化レベル "
        "（なし／軽量／標準／高圧縮／最大圧縮／カスタム）と「テキストデータを削除（画像のみ）」を選べます。"
        "PNG・JPEG の場合は DPI（72〜600）と JPEG 品質を指定できます。",
    )
    _img(doc, "10_export_dialog.png", width_cm=10)

    _h2(doc, "3.14　ファイルを外に出す")
    _para(
        doc,
        "メイン画面のカードを JusticePDF の外（エクスプローラーや他のアプリ）へドラッグ＆ドロップすれば、"
        "ファイルとしてそのまま取り出せます。右クリックメニューからは「印刷」や「他のアプリで開く」も選べます。",
    )

    _h2(doc, "3.15　マルチウィンドウ")
    _para(
        doc,
        "フォルダカードをダブルクリックすると、そのフォルダを別ウィンドウで開けます。"
        "別の作業フォルダを並べて、ドラッグ＆ドロップでファイルやフォルダをやり取りできます。",
    )
    _img(doc, "12_multi_window.png")

    return doc


def main() -> None:
    doc = build_doc()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
